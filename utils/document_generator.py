"""
Generador de Documentos PDF y DOCX
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import Dict


class DocumentGenerator:
    """Generador de documentos PDF y DOCX"""
    
    def __init__(self):
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generar_texto_vista_previa(self, datos: Dict) -> str:
        """Genera texto plano para vista previa"""
        texto = "\n"
        texto += "COLEGIO DE MÉDICOS\n"
        texto += "DE LA PROVINCIA DE BUENOS AIRES\n"
        texto += "CONSEJO SUPERIOR\n"
        texto += "_" * 80 + "\n\n"
        
        texto += "REUNIÓN ORDINARIA\n"
        texto += "ORDEN DEL DÍA\n\n"
        
        texto += f"FECHA: {datos['fecha']}\n"
        texto += f"HORA: {datos['hora']}\n"
        texto += f"LUGAR: {datos['lugar']}\n"
        if datos.get('sede') and datos['sede'].strip():
            texto += f"SEDE: {datos['sede']}\n"
        if datos.get('tipo') == 'virtual' and datos.get('plataforma') and datos['plataforma'].strip():
            texto += f"PLATAFORMA: {datos['plataforma']}\n"
        texto += "\n"
        
        texto += "DELEGADOS TITULARES:\n"
        texto += "-" * 80 + "\n"
        for d in datos['delegados']:
            texto += f"{d['titulo']} {d['nombre']} {d['apellido']:<40} {d['distrito']}\n"
        
        texto += "\n" + "=" * 80 + "\n"
        texto += "ORDEN DEL DÍA\n"
        texto += "=" * 80 + "\n\n"
        
        for tema in datos['orden_dia']:
            texto += f"{tema['numero_orden']}.- {tema['descripcion']}\n\n"
        
        texto += "\n" + "-" * 80 + "\n"
        texto += "Saludamos a Ud. atentamente.\n\n"
        texto += f"Secretario General: {datos['secretario']}\n"
        texto += f"Presidente: {datos['presidente']}\n"
        texto += "-" * 80 + "\n"
        
        return texto
    
    def generar_pdf(self, datos: Dict) -> str:
        """Genera documento PDF con diseño profesional"""
        # Nombre de archivo
        fecha_archivo = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ORDEN_DEL_DIA_{fecha_archivo}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        # Crear documento
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=1.5*cm,
            leftMargin=1.5*cm,
            topMargin=1.5*cm,
            bottomMargin=1.5*cm
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Obtener tamaño y fuente del título desde datos
        tamaño_titulo = datos.get('tamaño_titulo', 12)
        fuente_titulo = datos.get('fuente_titulo', 'Helvetica')
        negrita_titulo = datos.get('negrita_titulo', True)
        negrita_subtitulo = datos.get('negrita_subtitulo', True)
        
        # Mapeo de fuentes compatibles con ReportLab
        font_map_reportlab = {
            'Helvetica': 'Helvetica-Bold',
            'Arial': 'Helvetica-Bold',
            'Times New Roman': 'Times-Bold',
            'Courier': 'Courier-Bold',
            'Georgia': 'Helvetica-Bold'
        }
        font_normal_map = {
            'Helvetica': 'Helvetica',
            'Arial': 'Helvetica',
            'Times New Roman': 'Times-Roman',
            'Courier': 'Courier',
            'Georgia': 'Helvetica'
        }
        
        # Seleccionar fuente según negrita
        fuente_titulo_final = font_map_reportlab.get(fuente_titulo, 'Helvetica-Bold') if negrita_titulo else font_normal_map.get(fuente_titulo, 'Helvetica')
        fuente_subtitulo_final = font_map_reportlab.get(fuente_titulo, 'Helvetica-Bold') if negrita_subtitulo else font_normal_map.get(fuente_titulo, 'Helvetica')
        
        style_organismo = ParagraphStyle(
            'Organismo',
            parent=styles['Heading1'],
            fontSize=9,
            textColor=colors.HexColor('#2E7D32'),
            spaceAfter=2,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        style_sede = ParagraphStyle(
            'Sede',
            parent=styles['Heading1'],
            fontSize=8,
            textColor=colors.HexColor('#2E7D32'),
            spaceAfter=8,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        style_title = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=tamaño_titulo,
            textColor=colors.black,
            spaceAfter=2,
            alignment=TA_CENTER,
            fontName=fuente_titulo_final
        )
        
        style_subtitle = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=tamaño_titulo,
            textColor=colors.black,
            spaceAfter=10,
            alignment=TA_CENTER,
            fontName=fuente_subtitulo_final
        )
        
        style_heading = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading3'],
            fontSize=10,
            textColor=colors.black,
            spaceAfter=6,
            spaceBefore=10,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            underline=True
        )
        
        style_normal = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica'
        )
        
        style_tema = ParagraphStyle(
            'CustomTema',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=8,
            leftIndent=0.2*cm,
            fontName='Helvetica'
        )
        
        # Contenido
        story = []
        
        # Línea decorativa superior
        story.append(Spacer(1, 0.2*cm))
        
        # Agregar logo si existe
        if datos.get('imagen_logo') and os.path.exists(datos['imagen_logo']):
            try:
                from reportlab.platypus import Image as RLImage
                # Obtener tamaño personalizado del logo o usar default
                ancho_logo = datos.get('ancho_logo', 3.5)  # cm
                alto_logo = datos.get('alto_logo', 2)  # cm
                logo = RLImage(datos['imagen_logo'], width=ancho_logo*cm, height=alto_logo*cm)
                logo_table = Table([[logo]], colWidths=[ancho_logo*cm])
                logo_table.setStyle(TableStyle([('ALIGN', (0, 0), (0, 0), 'CENTER')]))
                story.append(logo_table)
                story.append(Spacer(1, 0.3*cm))
            except Exception as e:
                print(f"Error cargando logo en PDF: {e}")
        
        # Título documento (usar texto personalizado)
        texto_encabezado = datos.get('texto_encabezado', 'ORDEN DEL DÍA')
        story.append(Paragraph(texto_encabezado, style_title))
        
        # Subtítulo si existe
        subtitulo = datos.get('subtitulo_encabezado', '').strip()
        if subtitulo:
            story.append(Paragraph(subtitulo, style_subtitle))
        
        story.append(Spacer(1, 0.3*cm))
        
        # Datos de la reunión
        datos_tabla = [
            ['FECHA:', datos['fecha']],
            ['HORA:', datos['hora']],
            ['LUGAR:', datos['lugar']]
        ]
        
        # Agregar SEDE solo si no está vacío
        if datos.get('sede') and datos['sede'].strip():
            datos_tabla.append(['SEDE:', datos['sede']])
        
        # Agregar PLATAFORMA si es reunión virtual
        if datos.get('tipo') == 'virtual' and datos.get('plataforma') and datos['plataforma'].strip():
            datos_tabla.append(['PLATAFORMA:', datos['plataforma']])
        
        table_datos = Table(datos_tabla, colWidths=[2*cm, 11*cm])
        table_datos.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        story.append(table_datos)
        story.append(Spacer(1, 0.4*cm))
        
        # Delegados titulares
        story.append(Paragraph("DELEGADOS TITULARES:", style_heading))
        
        # Tabla de delegados
        delegados_data = [['Nombre y Apellido', 'Distrito']]
        for d in datos['delegados']:
            delegados_data.append([
                f"{d['titulo']} {d['nombre']} {d['apellido']}",
                d['distrito']
            ])
        
        table_delegados = Table(delegados_data, colWidths=[9.5*cm, 3.5*cm])
        table_delegados.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
        ]))
        
        story.append(table_delegados)
        story.append(Spacer(1, 0.4*cm))
        
        # Salto de página
        story.append(PageBreak())
        
        # Agregar logo en la segunda página si existe
        if datos.get('imagen_logo') and os.path.exists(datos['imagen_logo']):
            try:
                ancho_logo = datos.get('ancho_logo', 3.5)  # cm
                alto_logo = datos.get('alto_logo', 2.0)    # cm
                img = Image(datos['imagen_logo'], width=ancho_logo*cm, height=alto_logo*cm)
                story.append(img)
                story.append(Spacer(1, 0.3*cm))
            except Exception as e:
                print(f"Error cargando logo en PDF: {e}")
        
        # Orden del día
        story.append(Paragraph("ORDEN DEL DÍA", style_heading))
        
        for tema in datos['orden_dia']:
            texto_tema = f"<b>{tema['numero_orden']}.-</b> {tema['descripcion']}"
            story.append(Paragraph(texto_tema, style_tema))
        
        story.append(Spacer(1, 0.6*cm))
        
        # Saludo
        story.append(Paragraph("Saludamos a Ud. atentamente.", style_normal))
        story.append(Spacer(1, 1.5*cm))
        
        # Firmas
        firmas_data = [
            [datos['secretario'], datos['presidente']],
            ['Secretario General', 'Presidente']
        ]
        
        table_firmas = Table(firmas_data, colWidths=[6.5*cm, 6.5*cm])
        table_firmas.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, 1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
            ('LINEABOVE', (0, 0), (-1, 0), 1, colors.black),
            ('TOPPADDING', (0, 0), (-1, 0), 30),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 2),
        ]))
        
        story.append(table_firmas)
        
        # Generar PDF
        doc.build(story)
        
        return filepath
    
    def generar_docx(self, datos: Dict) -> str:
        """Genera documento DOCX con diseño profesional"""
        # Nombre de archivo
        fecha_archivo = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ORDEN_DEL_DIA_{fecha_archivo}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Agregar logo si existe
        if datos.get('imagen_logo') and os.path.exists(datos['imagen_logo']):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                # Obtener tamaño personalizado del logo o usar default
                ancho_logo = datos.get('ancho_logo_docx', 1.2)  # pulgadas
                logo_run.add_picture(datos['imagen_logo'], width=Inches(ancho_logo))
                doc.add_paragraph()  # Espacio
            except Exception as e:
                print(f"Error cargando logo en DOCX: {e}")
        
        # Obtener tamaño y fuente del título desde datos
        tamaño_titulo = datos.get('tamaño_titulo', 12)
        fuente_titulo = datos.get('fuente_titulo', 'Helvetica')
        negrita_titulo = datos.get('negrita_titulo', True)
        negrita_subtitulo = datos.get('negrita_subtitulo', True)
        
        # Mapeo de fuentes a Windows
        font_map = {
            'Helvetica': 'Calibri',
            'Arial': 'Arial',
            'Times New Roman': 'Times New Roman',
            'Courier': 'Courier New',
            'Georgia': 'Georgia'
        }
        fuente_word = font_map.get(fuente_titulo, 'Calibri')
        
        # Título documento (usar texto personalizado)
        texto_encabezado = datos.get('texto_encabezado', 'ORDEN DEL DÍA')
        order_title = doc.add_paragraph(texto_encabezado)
        order_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in order_title.runs:
            run.font.bold = negrita_titulo
            run.font.size = Pt(tamaño_titulo)
            run.font.name = fuente_word
        
        # Subtítulo si existe
        subtitulo = datos.get('subtitulo_encabezado', '').strip()
        if subtitulo:
            subtitle = doc.add_paragraph(subtitulo)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.font.bold = negrita_subtitulo
                run.font.size = Pt(tamaño_titulo)
                run.font.name = fuente_word
        
        doc.add_paragraph()  # Espacio
        
        # Datos de la reunión
        p_fecha = doc.add_paragraph()
        p_fecha.add_run('FECHA: ').bold = True
        p_fecha.add_run(datos['fecha'])
        for run in p_fecha.runs:
            run.font.size = Pt(9)
        
        p_hora = doc.add_paragraph()
        p_hora.add_run('HORA: ').bold = True
        p_hora.add_run(datos['hora'])
        for run in p_hora.runs:
            run.font.size = Pt(9)
        
        p_lugar = doc.add_paragraph()
        p_lugar.add_run('LUGAR: ').bold = True
        p_lugar.add_run(datos['lugar'])
        for run in p_lugar.runs:
            run.font.size = Pt(9)
        
        # Agregar SEDE solo si no está vacío
        if datos.get('sede') and datos['sede'].strip():
            p_sede = doc.add_paragraph()
            p_sede.add_run('SEDE: ').bold = True
            p_sede.add_run(datos['sede'])
            for run in p_sede.runs:
                run.font.size = Pt(9)
        
        # Agregar PLATAFORMA si es reunión virtual
        if datos.get('tipo') == 'virtual' and datos.get('plataforma') and datos['plataforma'].strip():
            p_plataforma = doc.add_paragraph()
            p_plataforma.add_run('PLATAFORMA: ').bold = True
            p_plataforma.add_run(datos['plataforma'])
            for run in p_plataforma.runs:
                run.font.size = Pt(9)
        
        doc.add_paragraph()  # Espacio
        
        # Delegados titulares
        heading_del = doc.add_heading('DELEGADOS TITULARES:', 2)
        for run in heading_del.runs:
            run.font.size = Pt(10)
        
        # Tabla de delegados
        table = doc.add_table(rows=len(datos['delegados']) + 1, cols=2)
        table.style = 'Light Grid'
        
        # Encabezados tabla
        table.rows[0].cells[0].text = 'Nombre y Apellido'
        table.rows[0].cells[1].text = 'Distrito'
        
        # Formatear encabezados
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        # Datos delegados
        for i, d in enumerate(datos['delegados'], 1):
            nombre_completo = f"{d['titulo']} {d['nombre']} {d['apellido']}"
            table.rows[i].cells[0].text = nombre_completo
            table.rows[i].cells[1].text = d['distrito']
            
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        doc.add_paragraph()  # Espacio
        
        # Salto de página
        doc.add_page_break()
        
        # Agregar logo en la segunda página si existe
        if datos.get('imagen_logo') and os.path.exists(datos['imagen_logo']):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                ancho_logo = datos.get('ancho_logo_docx', 1.2)  # pulgadas
                logo_run.add_picture(datos['imagen_logo'], width=Inches(ancho_logo))
                doc.add_paragraph()  # Espacio
            except Exception as e:
                print(f"Error cargando logo en DOCX página 2: {e}")
        
        # Orden del día
        heading_orden = doc.add_heading('ORDEN DEL DÍA', 2)
        heading_orden.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading_orden.runs:
            run.font.size = Pt(10)
            run.underline = True
        
        doc.add_paragraph()  # Espacio pequeño
        
        for tema in datos['orden_dia']:
            p = doc.add_paragraph(f"{tema['numero_orden']}.- {tema['descripcion']}")
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(9)
        
        doc.add_paragraph()  # Espacio
        
        # Saludo
        saludo = doc.add_paragraph('Saludamos a Ud. atentamente.')
        for run in saludo.runs:
            run.font.size = Pt(9)
        
        doc.add_paragraph()  # Espacio
        doc.add_paragraph()  # Espacio
        doc.add_paragraph()  # Espacio para firmas
        
        # Firmas
        table_firmas = doc.add_table(rows=2, cols=2)
        table_firmas.style = 'Table Grid'
        
        # Primera fila: nombres
        table_firmas.rows[0].cells[0].text = datos['secretario']
        table_firmas.rows[0].cells[1].text = datos['presidente']
        
        # Segunda fila: cargos
        table_firmas.rows[1].cells[0].text = 'Secretario General'
        table_firmas.rows[1].cells[1].text = 'Presidente'
        
        # Centrar, formatear y agregar línea a firmas
        for row_idx, row in enumerate(table_firmas.rows):
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
                    if row_idx == 1:  # Cargos en tamaño menor
                        run.font.size = Pt(8)
        
        # Guardar
        doc.save(filepath)
        
        return filepath